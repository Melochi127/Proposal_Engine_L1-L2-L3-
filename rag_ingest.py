"""
Telecom Proposal Engine - Hybrid RAG Ingestion (FIXED)
======================================================
Fixes:
  1. .doc files now convert to .docx via LibreOffice/python-docx fallback
  2. Chunk size reduced from 800→500 for better retrieval precision
  3. Added metadata enrichment (source, page, section detection)
  4. Handles bm25_chunks.pkl in rag_data without crashing

Builds:
  1. Chroma vector DB (MiniLM embeddings)
  2. BM25 chunk store
"""

import os
import sys
import pickle
import shutil
import subprocess
import tempfile
from pathlib import Path
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from config import CHROMA_DIR, BM25_STORE, EMBEDDING_MODEL, RAG_DATA_DIR

# ─── Override chunk size for more chunks ─────────────────
# Change these to control chunk count:
#   500 chars → ~140-180 chunks from your 28 pages + .doc files
#   800 chars → ~67-90 chunks (your current problem)
CHUNK_SIZE = 500       # was 800 — smaller = more precise retrieval
CHUNK_OVERLAP = 100    # was 150


def build_hybrid_store(documents: list, reset_db: bool = True):
    """Build ChromaDB + BM25 from LangChain Document list."""

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(documents)

    # Enrich metadata
    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_id"] = i
        chunk.metadata["chunk_size"] = len(chunk.page_content)
        # Detect section headings in content
        lines = chunk.page_content.strip().split("\n")
        if lines and len(lines[0]) < 80 and lines[0].strip():
            chunk.metadata["section_heading"] = lines[0].strip()

    os.makedirs(os.path.dirname(BM25_STORE), exist_ok=True)

    # Save chunks for BM25
    with open(BM25_STORE, "wb") as f:
        pickle.dump(chunks, f)

    # Rebuild Chroma if requested
    if reset_db and os.path.exists(CHROMA_DIR):
        shutil.rmtree(CHROMA_DIR)

    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)

    vectordb = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=CHROMA_DIR,
    )

    return {
        "num_chunks": len(chunks),
        "chroma_dir": CHROMA_DIR,
        "bm25_store": BM25_STORE,
    }


# ─── .DOC file handling ──────────────────────────────────

def convert_doc_to_text(fpath: str) -> str:
    """
    Extract text from .doc files using multiple fallback methods.
    Priority: antiword → textract → python-docx (if renamed) → raw read
    """
    text = ""

    # Method 1: Try antiword (Linux) or textract
    for cmd in ["antiword", "catdoc"]:
        try:
            result = subprocess.run(
                [cmd, fpath],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0 and len(result.stdout.strip()) > 50:
                text = result.stdout
                print(f"    Extracted via {cmd}: {len(text)} chars")
                return text
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    # Method 2: Try LibreOffice headless conversion
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                ["soffice", "--headless", "--convert-to", "txt:Text", "--outdir", tmpdir, fpath],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode == 0:
                txt_files = [f for f in os.listdir(tmpdir) if f.endswith(".txt")]
                if txt_files:
                    txt_path = os.path.join(tmpdir, txt_files[0])
                    with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
                        text = f.read()
                    if len(text.strip()) > 50:
                        print(f"    Extracted via LibreOffice: {len(text)} chars")
                        return text
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Method 3: Try python-docx (works if .doc is actually a .docx internally)
    try:
        import docx
        doc = docx.Document(fpath)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if len(text.strip()) > 50:
            print(f"    Extracted via python-docx: {len(text)} chars")
            return text
    except Exception:
        pass

    # Method 4: Try docx2txt
    try:
        import docx2txt
        text = docx2txt.process(fpath)
        if text and len(text.strip()) > 50:
            print(f"    Extracted via docx2txt: {len(text)} chars")
            return text
    except Exception:
        pass

    # Method 5: Raw text extraction (last resort — binary .doc contains some readable text)
    try:
        with open(fpath, "rb") as f:
            raw = f.read()
        # Extract printable ASCII sequences > 20 chars
        import re
        strings = re.findall(rb"[\x20-\x7E]{20,}", raw)
        text = "\n".join(s.decode("ascii", errors="ignore") for s in strings)
        if len(text.strip()) > 100:
            print(f"    Extracted via raw binary scan: {len(text)} chars")
            return text
    except Exception:
        pass

    return ""


def load_doc_file(fpath: str, fname: str) -> list:
    """Load a .doc file and return list of LangChain Documents."""
    # First try UnstructuredWordDocumentLoader (if unstructured is installed)
    try:
        from langchain_community.document_loaders import UnstructuredWordDocumentLoader
        loaded = UnstructuredWordDocumentLoader(fpath).load()
        if loaded and any(len(d.page_content.strip()) > 50 for d in loaded):
            print(f"  Loaded DOC via Unstructured: {len(loaded)} sections")
            return loaded
    except Exception as e:
        print(f"    UnstructuredWordDocumentLoader failed: {e}")

    # Fallback: manual extraction
    text = convert_doc_to_text(fpath)
    if text and len(text.strip()) > 50:
        # Split into pages (~3000 chars each to simulate page breaks)
        page_size = 3000
        pages = []
        for i in range(0, len(text), page_size):
            page_text = text[i:i + page_size].strip()
            if page_text:
                pages.append(Document(
                    page_content=page_text,
                    metadata={"source": fname, "page": i // page_size},
                ))
        print(f"  Loaded DOC via text extraction: {len(pages)} pages")
        return pages

    print(f"  ⚠️ Could not extract text from {fname}")
    print(f"     FIX: Open in Word → Save As → .docx format")
    return []


# ─── Main ingestion ─────────────────────────────────────

def run_ingestion() -> bool:
    """Scan rag_data/ for documents, build hybrid BM25+Chroma stores."""
    from langchain_community.document_loaders import PyPDFLoader, TextLoader

    # Try importing docx loaders
    try:
        from langchain_community.document_loaders import Docx2txtLoader
        has_docx2txt = True
    except ImportError:
        has_docx2txt = False
        print("⚠️ docx2txt not installed. Install: pip install docx2txt")

    print(f"\n{'='*60}")
    print(f"  RAG INGESTION — Hybrid BM25 + ChromaDB")
    print(f"{'='*60}")
    print(f"  Data dir:    {RAG_DATA_DIR}")
    print(f"  Chunk size:  {CHUNK_SIZE} (overlap: {CHUNK_OVERLAP})")
    print(f"  Embedding:   {EMBEDDING_MODEL}")
    print(f"{'='*60}\n")

    if not os.path.exists(RAG_DATA_DIR):
        print("❌ rag_data folder does not exist.")
        return False

    docs = []
    files = sorted(os.listdir(RAG_DATA_DIR))
    doc_files = [f for f in files if os.path.isfile(os.path.join(RAG_DATA_DIR, f))]

    # Filter out non-document files
    skip_extensions = {".pkl", ".db", ".sqlite3", ".json", ".log", ".pyc"}
    skip_prefixes = {".", "__"}

    print(f"  Found {len(doc_files)} files\n")

    loaded_count = 0
    skipped_count = 0

    for fname in doc_files:
        fpath = os.path.join(RAG_DATA_DIR, fname)
        ext = os.path.splitext(fname)[1].lower()

        # Skip non-document files
        if ext in skip_extensions or any(fname.startswith(p) for p in skip_prefixes):
            print(f"  ⏭️  Skip: {fname} (not a document)")
            skipped_count += 1
            continue

        print(f"  📄 {fname}")

        try:
            if ext == ".pdf":
                loaded = PyPDFLoader(fpath).load()
                # Add source metadata
                for d in loaded:
                    d.metadata["source"] = fname
                docs.extend(loaded)
                loaded_count += 1
                print(f"     ✅ {len(loaded)} pages")

            elif ext == ".txt" or ext == ".md":
                loaded = TextLoader(fpath, encoding="utf-8").load()
                for d in loaded:
                    d.metadata["source"] = fname
                docs.extend(loaded)
                loaded_count += 1
                print(f"     ✅ {len(loaded)} sections")

            elif ext == ".docx":
                if has_docx2txt:
                    loaded = Docx2txtLoader(fpath).load()
                    for d in loaded:
                        d.metadata["source"] = fname
                    docs.extend(loaded)
                    loaded_count += 1
                    print(f"     ✅ {len(loaded)} sections")
                else:
                    # Fallback: try python-docx directly
                    try:
                        import docx
                        doc_obj = docx.Document(fpath)
                        text = "\n".join(p.text for p in doc_obj.paragraphs if p.text.strip())
                        if text.strip():
                            docs.append(Document(
                                page_content=text,
                                metadata={"source": fname, "page": 0},
                            ))
                            loaded_count += 1
                            print(f"     ✅ loaded via python-docx")
                        else:
                            print(f"     ⚠️ Empty document")
                    except Exception as e:
                        print(f"     ❌ Failed: {e}")
                        skipped_count += 1

            elif ext == ".doc":
                loaded = load_doc_file(fpath, fname)
                if loaded:
                    docs.extend(loaded)
                    loaded_count += 1
                else:
                    skipped_count += 1

            else:
                print(f"     ⏭️  Unsupported format")
                skipped_count += 1

        except Exception as e:
            print(f"     ❌ Error: {e}")
            skipped_count += 1

    print(f"\n{'─'*60}")
    print(f"  Loaded: {loaded_count} files → {len(docs)} document pages/sections")
    print(f"  Skipped: {skipped_count} files")
    print(f"{'─'*60}")

    if not docs:
        print("\n❌ No documents were loaded. Check:")
        print("   1. Are there PDF/DOCX/TXT files in rag_data/?")
        print("   2. For .doc files: convert to .docx in Word")
        print("   3. Install: pip install docx2txt python-docx pypdf")
        return False

    # Build stores
    print(f"\n  Building {CHUNK_SIZE}-char chunks...")
    result = build_hybrid_store(docs)

    print(f"\n{'='*60}")
    print(f"  ✅ INGESTION COMPLETE")
    print(f"{'='*60}")
    print(f"  Chunks built:  {result['num_chunks']}")
    print(f"  Chroma saved:  {result['chroma_dir']}")
    print(f"  BM25 saved:    {result['bm25_store']}")
    print(f"  Chunk size:    {CHUNK_SIZE} chars")

    # Warn if chunk count is low
    if result["num_chunks"] < 100:
        print(f"\n  ⚠️ Only {result['num_chunks']} chunks. To increase:")
        print(f"     - Reduce CHUNK_SIZE (currently {CHUNK_SIZE})")
        print(f"     - Add more documents to rag_data/")
        print(f"     - Convert .doc files to .docx")
    elif result["num_chunks"] >= 140:
        print(f"\n  ✅ Good chunk count for evaluation!")

    print()

    # Clear retriever cache if it exists
    try:
        from rag_retriever import clear_cache
        clear_cache()
        print("  Cache cleared.")
    except (ImportError, Exception):
        pass

    return True


if __name__ == "__main__":
    success = run_ingestion()
    if not success:
        sys.exit(1)